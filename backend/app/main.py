import asyncio
import re
from io import BytesIO
from urllib.parse import quote
from datetime import datetime, timedelta

import sqlite3

import httpx
from fastapi import FastAPI, HTTPException, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from .ai_client import generate_evening_feedback, generate_feedback, organize_lesson_note
from .ai_settings import (
    AIConfig,
    ai_settings_payload,
    delete_teacher_ai_config,
    decrypt_api_key,
    get_teacher_ai_config_row,
    require_teacher_ai_config,
    save_teacher_ai_config,
    save_teacher_ai_settings,
    set_active_ai_model,
    test_ai_connection,
)
from .database import get_db, init_db, now_iso
from .email_service import make_code, send_verification_email
from .schemas import (
    AISettingsTest,
    AISettingsUpdate,
    AIConfigCreate,
    AIConfigUpdate,
    AIModelSelection,
    EmailCodeRequest,
    EveningClassCreate,
    EveningClassUpdate,
    EveningFeedbackClassExportRequest,
    EveningFeedbackArchiveDeleteRequest,
    EveningFeedbackDeleteRequest,
    EveningFeedbackBatchExportRequest,
    EveningFeedbackBatchGenerateRequest,
    EveningFeedbackBatchSaveRequest,
    EveningFeedbackCreate,
    EveningFeedbackGenerateRequest,
    EveningFeedbackUpdate,
    EveningStudentBulkCreate,
    EveningStudentUpdate,
    FeedbackCreate,
    FeedbackDeleteRequest,
    FeedbackGenerateRequest,
    FeedbackOrganizeRequest,
    FeedbackOrganizeResponse,
    FeedbackUpdate,
    GroupClassCreate,
    GroupClassUpdate,
    LoginRequest,
    RegisterRequest,
    StudentCreate,
    StudentUpdate,
    StyleExampleCreate,
    StyleExampleFromFeedback,
    StyleExampleUpdate,
)
from .security import CurrentTeacher, create_token, hash_password, verify_password


app = FastAPI(title="教师一对一反馈助手 API")

MAX_ENABLED_STYLE_EXAMPLES = 5
EVENING_BATCH_GENERATE_CONCURRENCY = 50

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://127.0.0.1:5173", "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def on_startup() -> None:
    init_db()


def require_student(student_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        student = db.execute(
            "SELECT * FROM students WHERE id = ? AND teacher_id = ?",
            (student_id, teacher_id),
        ).fetchone()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    return dict(student)


def require_feedback(feedback_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        feedback = db.execute(
            "SELECT * FROM feedbacks WHERE id = ? AND teacher_id = ?",
            (feedback_id, teacher_id),
        ).fetchone()
    if not feedback:
        raise HTTPException(status_code=404, detail="反馈不存在")
    return dict(feedback)


def require_evening_class(class_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        row = db.execute(
            "SELECT * FROM evening_classes WHERE id = ? AND teacher_id = ?",
            (class_id, teacher_id),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="晚辅班级不存在")
    return dict(row)


def require_group_class(class_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        row = db.execute(
            "SELECT * FROM group_classes WHERE id = ? AND teacher_id = ?",
            (class_id, teacher_id),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="班课班级不存在")
    return dict(row)


def require_evening_student(student_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        row = db.execute(
            "SELECT * FROM evening_students WHERE id = ? AND teacher_id = ?",
            (student_id, teacher_id),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="晚辅学生不存在")
    return dict(row)


def require_evening_feedback(feedback_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        row = db.execute(
            "SELECT * FROM evening_feedbacks WHERE id = ? AND teacher_id = ?",
            (feedback_id, teacher_id),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="晚辅反馈不存在")
    return dict(row)


def require_evening_monthly_feedback(feedback_id: int, teacher_id: int) -> dict:
    return require_evening_feedback(feedback_id, teacher_id)


def evening_period_meta(period_type: str, period_value: str) -> dict[str, str]:
    try:
        if period_type == "day":
            value = datetime.strptime(period_value, "%Y-%m-%d").date()
            return {
                "period_type": "day",
                "period_value": period_value,
                "period_start": period_value,
                "period_end": period_value,
                "period_label": value.strftime("%Y-%m-%d"),
            }
        if period_type == "week":
            iso_year_text, iso_week_text = period_value.split("-W", 1)
            iso_year = int(iso_year_text)
            iso_week = int(iso_week_text)
            start = datetime.fromisocalendar(iso_year, iso_week, 1).date()
            end = start + timedelta(days=6)
            return {
                "period_type": "week",
                "period_value": f"{iso_year:04d}-W{iso_week:02d}",
                "period_start": start.isoformat(),
                "period_end": end.isoformat(),
                "period_label": f"{iso_year:04d}年第{iso_week:02d}周",
            }
        if period_type == "month":
            value = datetime.strptime(period_value, "%Y-%m").date()
            next_month = (value.replace(day=28) + timedelta(days=4)).replace(day=1)
            end = next_month - timedelta(days=1)
            return {
                "period_type": "month",
                "period_value": period_value,
                "period_start": value.isoformat(),
                "period_end": end.isoformat(),
                "period_label": value.strftime("%Y年%m月"),
            }
    except (TypeError, ValueError):
        pass
    raise HTTPException(status_code=400, detail="请选择有效的晚辅反馈时间")


def lesson_date_label(lesson_time: str) -> str:
    try:
        normalized = lesson_time.replace("Z", "").replace("T", " ")
        value = datetime.fromisoformat(normalized)
        return f"{value.month}.{value.day}"
    except ValueError:
        return lesson_time[:10] or "本次"


def lesson_number_from_title(title: str | None) -> int | None:
    match = re.search(r"第\s*(\d+)\s*次", title or "")
    return int(match.group(1)) if match else None


def next_lesson_number_from_feedbacks(feedbacks: list[dict]) -> int:
    recent_number = lesson_number_from_title(feedbacks[0]["lesson_title"]) if feedbacks else None
    if recent_number:
        return recent_number + 1
    max_number = max(
        (number for feedback in feedbacks if (number := lesson_number_from_title(feedback["lesson_title"]))),
        default=0,
    )
    return max_number + 1 if max_number else len(feedbacks) + 1


def safe_docx_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n]+', "", value).strip()
    cleaned = re.sub(r"\.docx$", "", cleaned, flags=re.IGNORECASE).strip()
    return cleaned or "晚辅反馈"


def export_period_filename_label(period: dict[str, str]) -> str:
    period_type = period.get("period_type")
    period_value = period.get("period_value", "")
    try:
        if period_type == "day":
            value = datetime.strptime(period_value, "%Y-%m-%d").date()
            return f"{value.month}.{value.day}"
        if period_type == "week":
            iso_year_text, iso_week_text = period_value.split("-W", 1)
            start = datetime.fromisocalendar(int(iso_year_text), int(iso_week_text), 1).date()
            return f"{start.month}月第{((start.day - 1) // 7) + 1}周"
        if period_type == "month":
            value = datetime.strptime(period_value, "%Y-%m").date()
            return f"{value.month}月"
    except (TypeError, ValueError):
        pass
    return period.get("period_label", period_value)


def evening_feedback_word_response(
    evening_class: dict,
    period: dict[str, str],
    term_label: str,
    owner_name: str,
    export_subject: str,
    document_title: str,
    filename_base: str,
    export_items: list[dict[str, str]],
) -> StreamingResponse:
    if not export_items:
        raise HTTPException(status_code=400, detail="没有可导出的晚辅反馈")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="服务器未安装 Word 导出依赖 python-docx") from exc

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(11)

    title_text = document_title.strip() or f"{term_label.strip()}{evening_class['name']}晚辅"
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(title_text.strip())
    title_run.bold = True
    title_run.font.name = "宋体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    title_run.font.size = Pt(18)

    document.add_paragraph()
    for item in export_items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        name_run = paragraph.add_run(f"{item['student_name']}：")
        name_run.bold = True
        name_run.font.name = "宋体"
        name_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        content_run = paragraph.add_run(item["final_feedback"])
        content_run.font.name = "宋体"
        content_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    suffix = f"{owner_name.strip()}{export_subject.strip()}".strip()
    generated_filename_base = safe_docx_filename(
        f"{term_label.strip()}{export_period_filename_label(period)}{evening_class['name']}晚辅反馈"
        f"{f'——{suffix}' if suffix else ''}"
    )
    filename = f"{safe_docx_filename(filename_base) if filename_base.strip() else generated_filename_base}.docx"
    encoded_filename = quote(filename)
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
    }
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


def normalize_style_feedback_type(feedback_type: str) -> str:
    return "evening_feedback" if feedback_type in ("evening_feedback", "evening_monthly") else "one_on_one"


def list_enabled_style_examples(teacher_id: int, feedback_type: str = "one_on_one", limit: int = 5) -> list[dict]:
    normalized_type = normalize_style_feedback_type(feedback_type)
    with get_db() as db:
        rows = db.execute(
            """
            SELECT id, title, content, feedback_type, source_type, source_feedback_id, enabled, created_at, updated_at
            FROM teacher_style_examples
            WHERE teacher_id = ? AND feedback_type = ? AND enabled = 1
            ORDER BY id DESC
            LIMIT ?
            """,
            (teacher_id, normalized_type, limit),
        ).fetchall()
    return [dict(row) for row in rows]


def count_enabled_style_examples(
    teacher_id: int,
    feedback_type: str = "one_on_one",
    exclude_example_id: int | None = None,
) -> int:
    normalized_type = normalize_style_feedback_type(feedback_type)
    query = """
        SELECT COUNT(*) AS count
        FROM teacher_style_examples
        WHERE teacher_id = ? AND feedback_type = ? AND enabled = 1
    """
    params: list[int | str] = [teacher_id, normalized_type]
    if exclude_example_id is not None:
        query += " AND id != ?"
        params.append(exclude_example_id)
    with get_db() as db:
        return db.execute(query, params).fetchone()["count"]


def ensure_style_example_enable_slot(
    teacher_id: int,
    feedback_type: str = "one_on_one",
    exclude_example_id: int | None = None,
) -> None:
    if count_enabled_style_examples(teacher_id, feedback_type, exclude_example_id) >= MAX_ENABLED_STYLE_EXAMPLES:
        raise HTTPException(status_code=400, detail="最多启用 5 条风格样例参与生成，建议 1-3 条；请先停用一条差异较大的样例")


def require_style_example(example_id: int, teacher_id: int) -> dict:
    with get_db() as db:
        row = db.execute(
            "SELECT * FROM teacher_style_examples WHERE id = ? AND teacher_id = ?",
            (example_id, teacher_id),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="风格样例不存在")
    return dict(row)


def _service_error_excerpt(text: str) -> str:
    compact = " ".join((text or "").split())
    return compact[:300]


def ai_error_detail(exc: Exception, action: str) -> str:
    if isinstance(exc, httpx.HTTPStatusError):
        status_code = exc.response.status_code
        excerpt = _service_error_excerpt(exc.response.text)
        suffix = f" 服务商返回：{excerpt}" if excerpt else ""
        if status_code in (401, 403):
            return f"{action}失败：API Key 无效、权限不足或账号无权访问该模型，请检查设置页里的 API Key 和模型权限。{suffix}"
        if status_code == 404:
            return f"{action}失败：接口地址或模型路径不存在，请检查 Base URL 是否包含正确的 /v1 兼容地址，以及模型名是否正确。{suffix}"
        if status_code == 429:
            return f"{action}失败：模型服务限流或额度不足，请稍后重试，或检查账号余额/调用额度。{suffix}"
        if status_code == 400:
            return f"{action}失败：请求参数、模型名或模型能力不匹配，请检查模型名和当前模型是否支持本功能。{suffix}"
        if status_code >= 500:
            return f"{action}失败：模型服务商暂时异常，请稍后重试。{suffix}"
        return f"{action}失败：模型服务返回 HTTP {status_code}。{suffix}"

    if isinstance(exc, httpx.TimeoutException):
        return f"{action}失败：连接模型服务超时，请稍后重试，或检查 Base URL 和服务商状态。"
    if isinstance(exc, httpx.RequestError):
        return f"{action}失败：无法连接模型服务，请检查 Base URL、网络连接或服务商状态。错误：{exc}"
    if isinstance(exc, (KeyError, IndexError, TypeError, ValueError)):
        return f"{action}失败：AI 返回内容格式异常，不符合 OpenAI-compatible Chat Completions 格式，请检查 Base URL 和模型接口。"
    return f"{action}失败：发生未知错误，请检查模型配置后重试。"


def raise_ai_exception(exc: Exception, action: str) -> None:
    raise HTTPException(status_code=502, detail=ai_error_detail(exc, action)) from exc


@app.get("/api/health")
def health():
    return {"ok": True}


@app.post("/api/auth/send-code")
def send_code(payload: EmailCodeRequest):
    code = make_code()
    expires_at = (datetime.utcnow() + timedelta(minutes=10)).isoformat(timespec="seconds") + "Z"
    with get_db() as db:
        db.execute(
            "INSERT INTO verification_codes (email, code, expires_at, used, created_at) VALUES (?, ?, ?, 0, ?)",
            (payload.email, code, expires_at, now_iso()),
        )
    sent = send_verification_email(payload.email, code)
    return {"message": "验证码已发送" if sent else "开发模式：验证码已打印在后端终端"}


@app.post("/api/auth/register")
def register(payload: RegisterRequest):
    with get_db() as db:
        exists = db.execute("SELECT id FROM teachers WHERE email = ?", (payload.email,)).fetchone()
        if exists:
            raise HTTPException(status_code=400, detail="该邮箱已注册")

        code = db.execute(
            """
            SELECT id, expires_at FROM verification_codes
            WHERE email = ? AND code = ? AND used = 0
            ORDER BY id DESC LIMIT 1
            """,
            (payload.email, payload.code),
        ).fetchone()
        if not code:
            raise HTTPException(status_code=400, detail="验证码不正确")
        if datetime.fromisoformat(code["expires_at"].removesuffix("Z")) < datetime.utcnow():
            raise HTTPException(status_code=400, detail="验证码已过期")

        cursor = db.execute(
            "INSERT INTO teachers (email, password_hash, created_at) VALUES (?, ?, ?)",
            (payload.email, hash_password(payload.password), now_iso()),
        )
        db.execute("UPDATE verification_codes SET used = 1 WHERE id = ?", (code["id"],))
        teacher_id = cursor.lastrowid

    return {"token": create_token(teacher_id), "teacher": {"id": teacher_id, "email": payload.email}}


@app.post("/api/auth/login")
def login(payload: LoginRequest):
    with get_db() as db:
        teacher = db.execute(
            "SELECT id, email, password_hash FROM teachers WHERE email = ?",
            (payload.email,),
        ).fetchone()
    if not teacher or not verify_password(payload.password, teacher["password_hash"]):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="邮箱或密码错误")
    return {"token": create_token(teacher["id"]), "teacher": {"id": teacher["id"], "email": teacher["email"]}}


@app.get("/api/auth/me")
def me(teacher: dict = CurrentTeacher):
    return {"teacher": teacher}


@app.get("/api/settings/ai")
def get_ai_settings(teacher: dict = CurrentTeacher):
    return ai_settings_payload(teacher["id"])


@app.put("/api/settings/ai")
def update_ai_settings(payload: AISettingsUpdate, teacher: dict = CurrentTeacher):
    row = save_teacher_ai_settings(
        teacher_id=teacher["id"],
        provider=payload.provider,
        base_url=payload.base_url,
        model=payload.model,
        api_key=payload.api_key,
        clear_api_key=payload.clear_api_key,
        feedback_format_mode=payload.feedback_format_mode,
    )
    return ai_settings_payload(teacher["id"])


@app.post("/api/settings/ai/configs")
def create_ai_config(payload: AIConfigCreate, teacher: dict = CurrentTeacher):
    save_teacher_ai_config(
        teacher_id=teacher["id"],
        name=payload.name,
        provider=payload.provider,
        base_url=payload.base_url,
        model=payload.model,
        api_key=payload.api_key,
        make_active=payload.make_active,
    )
    return ai_settings_payload(teacher["id"])


@app.put("/api/settings/ai/configs/{config_id}")
def update_ai_config(config_id: int, payload: AIConfigUpdate, teacher: dict = CurrentTeacher):
    save_teacher_ai_config(
        teacher_id=teacher["id"],
        config_id=config_id,
        name=payload.name,
        provider=payload.provider,
        base_url=payload.base_url,
        model=payload.model,
        api_key=payload.api_key,
        clear_api_key=payload.clear_api_key,
        make_active=payload.make_active,
    )
    return ai_settings_payload(teacher["id"])


@app.delete("/api/settings/ai/configs/{config_id}")
def delete_ai_config(config_id: int, teacher: dict = CurrentTeacher):
    delete_teacher_ai_config(teacher["id"], config_id)
    return ai_settings_payload(teacher["id"])


@app.post("/api/settings/ai/select")
def select_ai_model(payload: AIModelSelection, teacher: dict = CurrentTeacher):
    set_active_ai_model(
        teacher["id"],
        payload.model_type,
        config_id=payload.config_id,
        platform_model_id=payload.platform_model_id,
    )
    return ai_settings_payload(teacher["id"])


@app.post("/api/settings/ai/test")
async def test_ai_settings(payload: AISettingsTest, teacher: dict = CurrentTeacher):
    api_key = payload.api_key.strip()
    if not api_key:
        if payload.config_id:
            saved = get_teacher_ai_config_row(teacher["id"], payload.config_id)
            saved_key = decrypt_api_key(saved["encrypted_api_key"]) if saved else ""
            if not saved_key:
                raise HTTPException(status_code=400, detail="这条模型配置还没有可用的 API Key")
            config = AIConfig(
                api_key=saved_key,
                base_url=payload.base_url.rstrip("/"),
                model=payload.model,
                provider=payload.provider,
                source="personal",
                config_id=payload.config_id,
            )
        else:
            raise HTTPException(status_code=400, detail="请先填写 API Key，或先保存一份可用配置")
    else:
        config = AIConfig(
            api_key=api_key,
            base_url=payload.base_url.rstrip("/"),
            model=payload.model,
            provider=payload.provider,
        )
    reply = await test_ai_connection(config)
    return {"ok": True, "message": "连接成功，可以生成反馈", "reply": reply}


@app.get("/api/settings/style-examples")
def list_style_examples(feedback_type: str = "", teacher: dict = CurrentTeacher):
    normalized_type = normalize_style_feedback_type(feedback_type) if feedback_type else ""
    where_clause = "WHERE teacher_id = ?"
    params: list[int | str] = [teacher["id"]]
    if normalized_type:
        where_clause += " AND feedback_type = ?"
        params.append(normalized_type)
    with get_db() as db:
        rows = db.execute(
            f"""
            SELECT id, title, content, feedback_type, source_type, source_feedback_id, enabled, created_at, updated_at
            FROM teacher_style_examples
            {where_clause}
            ORDER BY id DESC
            """,
            params,
        ).fetchall()
    return {"examples": [dict(row) for row in rows]}


@app.post("/api/settings/style-examples")
def create_style_example(payload: StyleExampleCreate, teacher: dict = CurrentTeacher):
    feedback_type = normalize_style_feedback_type(payload.feedback_type)
    if payload.enabled:
        ensure_style_example_enable_slot(teacher["id"], feedback_type)
    timestamp = now_iso()
    with get_db() as db:
        cursor = db.execute(
            """
            INSERT INTO teacher_style_examples (
                teacher_id, title, content, feedback_type, source_type, source_feedback_id, enabled, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, 'manual', NULL, ?, ?, ?)
            """,
            (
                teacher["id"],
                payload.title.strip(),
                payload.content.strip(),
                feedback_type,
                int(payload.enabled),
                timestamp,
                timestamp,
            ),
        )
        row = db.execute("SELECT * FROM teacher_style_examples WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"example": dict(row)}


@app.put("/api/settings/style-examples/{example_id}")
def update_style_example(example_id: int, payload: StyleExampleUpdate, teacher: dict = CurrentTeacher):
    existing = require_style_example(example_id, teacher["id"])
    feedback_type = normalize_style_feedback_type(existing["feedback_type"])
    if payload.enabled and not existing["enabled"]:
        ensure_style_example_enable_slot(teacher["id"], feedback_type, example_id)
    with get_db() as db:
        db.execute(
            """
            UPDATE teacher_style_examples
            SET title = ?, content = ?, enabled = ?, updated_at = ?
            WHERE id = ? AND teacher_id = ?
            """,
            (
                payload.title.strip(),
                payload.content.strip(),
                int(payload.enabled),
                now_iso(),
                example_id,
                teacher["id"],
            ),
        )
        row = db.execute(
            "SELECT * FROM teacher_style_examples WHERE id = ? AND teacher_id = ?",
            (example_id, teacher["id"]),
        ).fetchone()
    return {"example": dict(row)}


@app.delete("/api/settings/style-examples/{example_id}")
def delete_style_example(example_id: int, teacher: dict = CurrentTeacher):
    require_style_example(example_id, teacher["id"])
    with get_db() as db:
        db.execute(
            "DELETE FROM teacher_style_examples WHERE id = ? AND teacher_id = ?",
            (example_id, teacher["id"]),
        )
    return {"ok": True}


@app.post("/api/settings/style-examples/from-feedback")
def create_style_example_from_feedback(payload: StyleExampleFromFeedback, teacher: dict = CurrentTeacher):
    if payload.feedback_type == "one_on_one":
        feedback = require_feedback(payload.feedback_id, teacher["id"])
        content = feedback["final_feedback"]
        title = payload.title.strip() or feedback["lesson_title"] or f"一对一反馈样例 {feedback['lesson_time']}"
        feedback_type = "one_on_one"
    else:
        feedback = require_evening_feedback(payload.feedback_id, teacher["id"])
        content = feedback["final_feedback"]
        title = payload.title.strip() or f"晚辅反馈样例 {feedback['period_label']}"
        feedback_type = "evening_feedback"

    ensure_style_example_enable_slot(teacher["id"], feedback_type)
    timestamp = now_iso()
    with get_db() as db:
        cursor = db.execute(
            """
            INSERT INTO teacher_style_examples (
                teacher_id, title, content, feedback_type, source_type, source_feedback_id, enabled, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, 'feedback', ?, 1, ?, ?)
            """,
            (teacher["id"], title, content, feedback_type, payload.feedback_id, timestamp, timestamp),
        )
        row = db.execute("SELECT * FROM teacher_style_examples WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"example": dict(row)}


@app.get("/api/students")
def list_students(teacher: dict = CurrentTeacher):
    with get_db() as db:
        rows = db.execute(
            """
            SELECT s.*, COUNT(f.id) AS feedback_count
            FROM students s
            LEFT JOIN feedbacks f ON f.student_id = s.id
            WHERE s.teacher_id = ?
            GROUP BY s.id
            ORDER BY s.id DESC
            """,
            (teacher["id"],),
        ).fetchall()
    return {"students": [dict(row) for row in rows]}


@app.post("/api/students")
def create_student(payload: StudentCreate, teacher: dict = CurrentTeacher):
    with get_db() as db:
        cursor = db.execute(
            """
            INSERT INTO students (teacher_id, name, grade, subject, note, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (teacher["id"], payload.name, payload.grade, payload.subject, payload.note, now_iso()),
        )
        student = db.execute("SELECT * FROM students WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"student": dict(student)}


@app.get("/api/students/{student_id}")
def get_student(student_id: int, teacher: dict = CurrentTeacher):
    return {"student": require_student(student_id, teacher["id"])}


@app.put("/api/students/{student_id}")
def update_student(student_id: int, payload: StudentUpdate, teacher: dict = CurrentTeacher):
    require_student(student_id, teacher["id"])
    with get_db() as db:
        db.execute(
            """
            UPDATE students
            SET name = ?, grade = ?, subject = ?, note = ?
            WHERE id = ? AND teacher_id = ?
            """,
            (payload.name, payload.grade, payload.subject, payload.note, student_id, teacher["id"]),
        )
        student = db.execute(
            "SELECT * FROM students WHERE id = ? AND teacher_id = ?",
            (student_id, teacher["id"]),
        ).fetchone()
    return {"student": dict(student)}


@app.delete("/api/students/{student_id}")
def delete_student(student_id: int, teacher: dict = CurrentTeacher):
    require_student(student_id, teacher["id"])
    with get_db() as db:
        db.execute("DELETE FROM students WHERE id = ? AND teacher_id = ?", (student_id, teacher["id"]))
    return {"ok": True}


@app.get("/api/students/{student_id}/feedbacks")
def list_feedbacks(student_id: int, teacher: dict = CurrentTeacher):
    require_student(student_id, teacher["id"])
    with get_db() as db:
        rows = db.execute(
            """
            SELECT * FROM feedbacks
            WHERE student_id = ? AND teacher_id = ?
            ORDER BY lesson_time DESC, id DESC
            """,
            (student_id, teacher["id"]),
        ).fetchall()
    return {"feedbacks": [dict(row) for row in rows]}


@app.post("/api/students/{student_id}/feedbacks/organize", response_model=FeedbackOrganizeResponse)
async def organize_feedback_note(student_id: int, payload: FeedbackOrganizeRequest, teacher: dict = CurrentTeacher):
    student = require_student(student_id, teacher["id"])
    if not (
        payload.raw_lesson_note.strip()
        or payload.lesson_summary.strip()
        or payload.performance_summary.strip()
        or payload.advice_summary.strip()
        or payload.homework_plan.strip()
    ):
        raise HTTPException(status_code=400, detail="请先填写本节课原始记录")
    ai_config = require_teacher_ai_config(
        teacher["id"],
        model_type=payload.model_type,
        config_id=payload.config_id,
        platform_model_id=payload.platform_model_id,
    )
    try:
        return await organize_lesson_note(
            student_name=student["name"],
            subject=student["subject"] or "数学",
            lesson_title=payload.lesson_title,
            raw_lesson_note=payload.raw_lesson_note,
            lesson_summary=payload.lesson_summary,
            performance_summary=payload.performance_summary,
            advice_summary=payload.advice_summary,
            homework_plan=payload.homework_plan,
            ai_config=ai_config,
        )
    except (httpx.HTTPError, KeyError, IndexError, TypeError, ValueError) as exc:
        raise_ai_exception(exc, "课堂记录整理")
    except Exception as exc:
        raise_ai_exception(exc, "课堂记录整理")


@app.post("/api/students/{student_id}/feedbacks/generate")
async def create_ai_draft(student_id: int, payload: FeedbackGenerateRequest, teacher: dict = CurrentTeacher):
    student = require_student(student_id, teacher["id"])
    with get_db() as db:
        feedback_rows = db.execute(
            """
            SELECT lesson_title
            FROM feedbacks
            WHERE student_id = ? AND teacher_id = ?
            ORDER BY lesson_time DESC, id DESC
            """,
            (student_id, teacher["id"]),
        ).fetchall()
    lesson_number = next_lesson_number_from_feedbacks([dict(row) for row in feedback_rows])
    ai_config = require_teacher_ai_config(
        teacher["id"],
        model_type=payload.model_type,
        config_id=payload.config_id,
        platform_model_id=payload.platform_model_id,
    )
    try:
        draft = await generate_feedback(
            student_name=student["name"],
            subject=student["subject"] or "数学",
            lesson_number=lesson_number,
            lesson_title=payload.lesson_title,
            lesson_date=lesson_date_label(payload.lesson_time),
            lesson_summary=payload.lesson_summary,
            performance_summary=payload.performance_summary,
            advice_summary=payload.advice_summary,
            homework_plan=payload.homework_plan,
            style_examples=list_enabled_style_examples(teacher["id"], "one_on_one") if payload.use_style_examples else [],
            ai_config=ai_config,
        )
    except (httpx.HTTPError, KeyError, IndexError, TypeError, ValueError) as exc:
        raise_ai_exception(exc, "AI 初稿生成")
    except Exception as exc:
        raise_ai_exception(exc, "AI 初稿生成")
    return {"draft": draft}


@app.post("/api/students/{student_id}/feedbacks")
def create_feedback(student_id: int, payload: FeedbackCreate, teacher: dict = CurrentTeacher):
    require_student(student_id, teacher["id"])
    timestamp = now_iso()
    with get_db() as db:
        cursor = db.execute(
            """
            INSERT INTO feedbacks (
                teacher_id, student_id, lesson_title, lesson_time, lesson_summary, performance_summary,
                advice_summary, homework_plan, ai_draft, final_feedback, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                teacher["id"],
                student_id,
                payload.lesson_title,
                payload.lesson_time,
                payload.lesson_summary,
                payload.performance_summary,
                payload.advice_summary,
                payload.homework_plan,
                payload.ai_draft,
                payload.final_feedback,
                timestamp,
                timestamp,
            ),
        )
        feedback = db.execute("SELECT * FROM feedbacks WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"feedback": dict(feedback)}


@app.get("/api/feedbacks")
def search_feedbacks(
    start_date: str | None = Query(default=None),
    end_date: str | None = Query(default=None),
    student_name: str = Query(default=""),
    teacher: dict = CurrentTeacher,
):
    conditions = ["f.teacher_id = ?"]
    params: list[str | int] = [teacher["id"]]

    if start_date:
        conditions.append("f.lesson_time >= ?")
        params.append(f"{start_date}T00:00")
    if end_date:
        conditions.append("f.lesson_time <= ?")
        params.append(f"{end_date}T23:59:59")
    if student_name.strip():
        conditions.append("s.name LIKE ?")
        params.append(f"%{student_name.strip()}%")

    with get_db() as db:
        rows = db.execute(
            f"""
            SELECT
                f.id, f.student_id, s.name AS student_name, s.grade, s.subject,
                f.lesson_title, f.lesson_time, f.lesson_summary, f.final_feedback,
                f.created_at, f.updated_at
            FROM feedbacks f
            JOIN students s ON s.id = f.student_id AND s.teacher_id = f.teacher_id
            WHERE {' AND '.join(conditions)}
            ORDER BY f.lesson_time DESC, f.id DESC
            """,
            params,
        ).fetchall()
    return {"feedbacks": [dict(row) for row in rows]}


@app.delete("/api/feedbacks/batch")
def delete_feedback_batch(payload: FeedbackDeleteRequest, teacher: dict = CurrentTeacher):
    ids = [feedback_id for feedback_id in payload.ids if feedback_id]
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要删除的一对一反馈")
    placeholders = ",".join("?" for _ in ids)
    with get_db() as db:
        cursor = db.execute(
            f"DELETE FROM feedbacks WHERE teacher_id = ? AND id IN ({placeholders})",
            (teacher["id"], *ids),
        )
    return {"ok": True, "deleted_count": cursor.rowcount or 0}


@app.get("/api/feedbacks/{feedback_id}")
def get_feedback(feedback_id: int, teacher: dict = CurrentTeacher):
    return {"feedback": require_feedback(feedback_id, teacher["id"])}


@app.put("/api/feedbacks/{feedback_id}")
def update_feedback(feedback_id: int, payload: FeedbackUpdate, teacher: dict = CurrentTeacher):
    require_feedback(feedback_id, teacher["id"])
    timestamp = now_iso()
    with get_db() as db:
        db.execute(
            """
            UPDATE feedbacks
            SET lesson_title = ?, lesson_time = ?, lesson_summary = ?, performance_summary = ?,
                advice_summary = ?, homework_plan = ?, ai_draft = ?, final_feedback = ?, updated_at = ?
            WHERE id = ? AND teacher_id = ?
            """,
            (
                payload.lesson_title,
                payload.lesson_time,
                payload.lesson_summary,
                payload.performance_summary,
                payload.advice_summary,
                payload.homework_plan,
                payload.ai_draft,
                payload.final_feedback,
                timestamp,
                feedback_id,
                teacher["id"],
            ),
        )
        feedback = db.execute(
            "SELECT * FROM feedbacks WHERE id = ? AND teacher_id = ?",
            (feedback_id, teacher["id"]),
        ).fetchone()
    return {"feedback": dict(feedback)}


@app.delete("/api/feedbacks/{feedback_id}")
def delete_feedback(feedback_id: int, teacher: dict = CurrentTeacher):
    require_feedback(feedback_id, teacher["id"])
    with get_db() as db:
        db.execute("DELETE FROM feedbacks WHERE id = ? AND teacher_id = ?", (feedback_id, teacher["id"]))
    return {"ok": True}


@app.get("/api/evening/classes")
def list_evening_classes(teacher: dict = CurrentTeacher):
    with get_db() as db:
        rows = db.execute(
            """
            SELECT c.*, COUNT(s.id) AS student_count
            FROM evening_classes c
            LEFT JOIN evening_students s ON s.class_id = c.id
            WHERE c.teacher_id = ?
            GROUP BY c.id
            ORDER BY c.id DESC
            """,
            (teacher["id"],),
        ).fetchall()
    return {"classes": [dict(row) for row in rows]}


@app.post("/api/evening/classes")
def create_evening_class(payload: EveningClassCreate, teacher: dict = CurrentTeacher):
    timestamp = now_iso()
    with get_db() as db:
        cursor = db.execute(
            """
            INSERT INTO evening_classes (teacher_id, name, created_at, updated_at)
            VALUES (?, ?, ?, ?)
            """,
            (teacher["id"], payload.name, timestamp, timestamp),
        )
        row = db.execute("SELECT * FROM evening_classes WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"class": dict(row)}


@app.get("/api/evening/classes/{class_id}")
def get_evening_class(class_id: int, teacher: dict = CurrentTeacher):
    return {"class": require_evening_class(class_id, teacher["id"])}


@app.put("/api/evening/classes/{class_id}")
def update_evening_class(class_id: int, payload: EveningClassUpdate, teacher: dict = CurrentTeacher):
    require_evening_class(class_id, teacher["id"])
    with get_db() as db:
        db.execute(
            "UPDATE evening_classes SET name = ?, updated_at = ? WHERE id = ? AND teacher_id = ?",
            (payload.name, now_iso(), class_id, teacher["id"]),
        )
        row = db.execute(
            "SELECT * FROM evening_classes WHERE id = ? AND teacher_id = ?",
            (class_id, teacher["id"]),
        ).fetchone()
    return {"class": dict(row)}


@app.delete("/api/evening/classes/{class_id}")
def delete_evening_class(class_id: int, teacher: dict = CurrentTeacher):
    require_evening_class(class_id, teacher["id"])
    with get_db() as db:
        db.execute("DELETE FROM evening_classes WHERE id = ? AND teacher_id = ?", (class_id, teacher["id"]))
    return {"ok": True}


@app.get("/api/group-classes")
def list_group_classes(teacher: dict = CurrentTeacher):
    with get_db() as db:
        rows = db.execute(
            """
            SELECT *
            FROM group_classes
            WHERE teacher_id = ?
            ORDER BY id DESC
            """,
            (teacher["id"],),
        ).fetchall()
    return {"classes": [dict(row) for row in rows]}


@app.post("/api/group-classes")
def create_group_class(payload: GroupClassCreate, teacher: dict = CurrentTeacher):
    timestamp = now_iso()
    with get_db() as db:
        cursor = db.execute(
            """
            INSERT INTO group_classes (teacher_id, name, created_at, updated_at)
            VALUES (?, ?, ?, ?)
            """,
            (teacher["id"], payload.name, timestamp, timestamp),
        )
        row = db.execute("SELECT * FROM group_classes WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"class": dict(row)}


@app.get("/api/group-classes/{class_id}")
def get_group_class(class_id: int, teacher: dict = CurrentTeacher):
    return {"class": require_group_class(class_id, teacher["id"])}


@app.put("/api/group-classes/{class_id}")
def update_group_class(class_id: int, payload: GroupClassUpdate, teacher: dict = CurrentTeacher):
    require_group_class(class_id, teacher["id"])
    with get_db() as db:
        db.execute(
            "UPDATE group_classes SET name = ?, updated_at = ? WHERE id = ? AND teacher_id = ?",
            (payload.name, now_iso(), class_id, teacher["id"]),
        )
        row = db.execute(
            "SELECT * FROM group_classes WHERE id = ? AND teacher_id = ?",
            (class_id, teacher["id"]),
        ).fetchone()
    return {"class": dict(row)}


@app.delete("/api/group-classes/{class_id}")
def delete_group_class(class_id: int, teacher: dict = CurrentTeacher):
    require_group_class(class_id, teacher["id"])
    with get_db() as db:
        db.execute("DELETE FROM group_classes WHERE id = ? AND teacher_id = ?", (class_id, teacher["id"]))
    return {"ok": True}


@app.get("/api/evening/classes/{class_id}/students")
def list_evening_students(class_id: int, teacher: dict = CurrentTeacher):
    require_evening_class(class_id, teacher["id"])
    with get_db() as db:
        rows = db.execute(
            """
            SELECT s.*, COUNT(f.id) AS feedback_count
            FROM evening_students s
            LEFT JOIN evening_feedbacks f ON f.student_id = s.id
            WHERE s.class_id = ? AND s.teacher_id = ?
            GROUP BY s.id
            ORDER BY s.id DESC
            """,
            (class_id, teacher["id"]),
        ).fetchall()
    return {"students": [dict(row) for row in rows]}


@app.post("/api/evening/classes/{class_id}/students/bulk")
def bulk_create_evening_students(class_id: int, payload: EveningStudentBulkCreate, teacher: dict = CurrentTeacher):
    require_evening_class(class_id, teacher["id"])
    names = []
    for line in payload.names_text.splitlines():
        name = line.strip()
        if name and name not in names:
            names.append(name)
    if not names:
        raise HTTPException(status_code=400, detail="请至少输入一名学生")

    timestamp = now_iso()
    with get_db() as db:
        for name in names:
            db.execute(
                """
                INSERT INTO evening_students (
                    teacher_id, class_id, name, grade, school, created_at, updated_at
                )
                VALUES (?, ?, ?, '', '', ?, ?)
                """,
                (teacher["id"], class_id, name, timestamp, timestamp),
            )
        rows = db.execute(
            "SELECT * FROM evening_students WHERE class_id = ? AND teacher_id = ? ORDER BY id DESC",
            (class_id, teacher["id"]),
        ).fetchall()
    return {"students": [dict(row) for row in rows], "created_count": len(names)}


@app.get("/api/evening/students/{student_id}")
def get_evening_student(student_id: int, teacher: dict = CurrentTeacher):
    return {"student": require_evening_student(student_id, teacher["id"])}


@app.put("/api/evening/students/{student_id}")
def update_evening_student(student_id: int, payload: EveningStudentUpdate, teacher: dict = CurrentTeacher):
    require_evening_student(student_id, teacher["id"])
    with get_db() as db:
        db.execute(
            """
            UPDATE evening_students
            SET name = ?, grade = ?, school = ?, updated_at = ?
            WHERE id = ? AND teacher_id = ?
            """,
            (payload.name, payload.grade, payload.school, now_iso(), student_id, teacher["id"]),
        )
        row = db.execute(
            "SELECT * FROM evening_students WHERE id = ? AND teacher_id = ?",
            (student_id, teacher["id"]),
        ).fetchone()
    return {"student": dict(row)}


@app.delete("/api/evening/students/{student_id}")
def delete_evening_student(student_id: int, teacher: dict = CurrentTeacher):
    require_evening_student(student_id, teacher["id"])
    with get_db() as db:
        db.execute("DELETE FROM evening_students WHERE id = ? AND teacher_id = ?", (student_id, teacher["id"]))
    return {"ok": True}


@app.get("/api/evening/students/{student_id}/feedbacks")
def list_evening_feedbacks(student_id: int, teacher: dict = CurrentTeacher):
    require_evening_student(student_id, teacher["id"])
    with get_db() as db:
        rows = db.execute(
            """
            SELECT * FROM evening_feedbacks
            WHERE student_id = ? AND teacher_id = ?
            ORDER BY period_start DESC, id DESC
            """,
            (student_id, teacher["id"]),
        ).fetchall()
    return {"feedbacks": [dict(row) for row in rows]}


@app.get("/api/evening/feedbacks")
def search_evening_feedbacks(
    start_date: str = Query(default=""),
    end_date: str = Query(default=""),
    period_type: str = Query(default=""),
    student_name: str = Query(default=""),
    class_id: int | None = Query(default=None),
    teacher: dict = CurrentTeacher,
):
    query = """
        SELECT f.*, s.name AS student_name, s.grade, s.school, c.name AS class_name
        FROM evening_feedbacks f
        JOIN evening_students s ON s.id = f.student_id
        JOIN evening_classes c ON c.id = s.class_id
        WHERE f.teacher_id = ?
    """
    params: list[str | int] = [teacher["id"]]
    if class_id:
        require_evening_class(class_id, teacher["id"])
        query += " AND c.id = ?"
        params.append(class_id)
    if period_type:
        if period_type not in {"day", "week", "month"}:
            raise HTTPException(status_code=400, detail="反馈类型无效")
        query += " AND f.period_type = ?"
        params.append(period_type)
    if student_name.strip():
        query += " AND s.name LIKE ?"
        params.append(f"%{student_name.strip()}%")
    if start_date:
        query += " AND f.period_end >= ?"
        params.append(start_date[:10])
    if end_date:
        query += " AND f.period_start <= ?"
        params.append(end_date[:10])
    query += " ORDER BY f.period_start DESC, f.id DESC"
    with get_db() as db:
        rows = db.execute(query, params).fetchall()
    return {"feedbacks": [dict(row) for row in rows]}


@app.get("/api/evening/feedbacks/archive")
def search_evening_feedback_archives(
    start_date: str = Query(default=""),
    end_date: str = Query(default=""),
    period_type: str = Query(default=""),
    class_id: int | None = Query(default=None),
    teacher: dict = CurrentTeacher,
):
    query = """
        SELECT
            c.id AS class_id,
            c.name AS class_name,
            f.period_type,
            f.period_value,
            f.period_start,
            f.period_end,
            f.period_label,
            COUNT(f.id) AS feedback_count,
            GROUP_CONCAT(DISTINCT NULLIF(f.subject, '')) AS subjects,
            MAX(f.id) AS latest_id,
            MAX(f.updated_at) AS updated_at
        FROM evening_feedbacks f
        JOIN evening_students s ON s.id = f.student_id
        JOIN evening_classes c ON c.id = s.class_id
        WHERE f.teacher_id = ?
    """
    params: list[str | int] = [teacher["id"]]
    if class_id:
        require_evening_class(class_id, teacher["id"])
        query += " AND c.id = ?"
        params.append(class_id)
    if period_type:
        if period_type not in {"day", "week", "month"}:
            raise HTTPException(status_code=400, detail="反馈类型无效")
        query += " AND f.period_type = ?"
        params.append(period_type)
    if start_date:
        query += " AND f.period_end >= ?"
        params.append(start_date[:10])
    if end_date:
        query += " AND f.period_start <= ?"
        params.append(end_date[:10])
    query += """
        GROUP BY c.id, c.name, f.period_type, f.period_value, f.period_start, f.period_end, f.period_label
        ORDER BY f.period_start DESC, latest_id DESC
    """
    with get_db() as db:
        rows = db.execute(query, params).fetchall()
    return {"archives": [dict(row) for row in rows]}


def evening_batch_feedback_row(db: sqlite3.Connection, feedback_id: int, teacher_id: int) -> dict:
    row = db.execute(
        """
        SELECT f.*, s.name AS student_name, s.grade, s.school, c.name AS class_name
        FROM evening_feedbacks f
        JOIN evening_students s ON s.id = f.student_id
        JOIN evening_classes c ON c.id = s.class_id
        WHERE f.id = ? AND f.teacher_id = ?
        """,
        (feedback_id, teacher_id),
    ).fetchone()
    return dict(row) if row else {}


@app.get("/api/evening/classes/{class_id}/feedbacks/batch")
def get_evening_feedback_batch(
    class_id: int,
    period_type: str = Query(...),
    period_value: str = Query(...),
    teacher: dict = CurrentTeacher,
):
    require_evening_class(class_id, teacher["id"])
    period = evening_period_meta(period_type, period_value)
    with get_db() as db:
        rows = db.execute(
            """
            SELECT
                s.id AS student_id,
                s.name AS student_name,
                s.grade,
                s.school,
                f.id AS feedback_id,
                f.subject,
                f.homework_summary,
                f.ai_draft,
                f.final_feedback,
                f.created_at AS feedback_created_at,
                f.updated_at AS feedback_updated_at
            FROM evening_students s
            LEFT JOIN evening_feedbacks f
                ON f.student_id = s.id
                AND f.teacher_id = ?
                AND f.period_type = ?
                AND f.period_value = ?
            WHERE s.class_id = ? AND s.teacher_id = ?
            ORDER BY s.id DESC
            """,
            (teacher["id"], period["period_type"], period["period_value"], class_id, teacher["id"]),
        ).fetchall()
    items = []
    for row in rows:
        record = dict(row)
        feedback = None
        if record["feedback_id"]:
            feedback = {
                "id": record["feedback_id"],
                "student_id": record["student_id"],
                "period_type": period["period_type"],
                "period_value": period["period_value"],
                "period_start": period["period_start"],
                "period_end": period["period_end"],
                "period_label": period["period_label"],
                "subject": record["subject"] or "",
                "homework_summary": record["homework_summary"] or "",
                "ai_draft": record["ai_draft"] or "",
                "final_feedback": record["final_feedback"] or "",
                "created_at": record["feedback_created_at"],
                "updated_at": record["feedback_updated_at"],
            }
        items.append(
            {
                "student": {
                    "id": record["student_id"],
                    "name": record["student_name"],
                    "grade": record["grade"],
                    "school": record["school"],
                    "class_id": class_id,
                },
                "feedback": feedback,
            }
        )
    return {"period": period, "items": items}


@app.get("/api/evening/classes/{class_id}/feedbacks/archive")
def list_evening_class_feedback_archive(class_id: int, teacher: dict = CurrentTeacher):
    require_evening_class(class_id, teacher["id"])
    with get_db() as db:
        rows = db.execute(
            """
            SELECT
                f.period_type,
                f.period_value,
                f.period_start,
                f.period_end,
                f.period_label,
                COUNT(f.id) AS feedback_count,
                GROUP_CONCAT(DISTINCT NULLIF(f.subject, '')) AS subjects,
                MAX(f.id) AS latest_id,
                MAX(f.updated_at) AS updated_at
            FROM evening_feedbacks f
            JOIN evening_students s ON s.id = f.student_id
            WHERE f.teacher_id = ? AND s.class_id = ?
            GROUP BY f.period_type, f.period_value, f.period_start, f.period_end, f.period_label
            ORDER BY f.period_start DESC, latest_id DESC
            """,
            (teacher["id"], class_id),
        ).fetchall()
    return {"archives": [dict(row) for row in rows]}


@app.post("/api/evening/classes/{class_id}/feedbacks/batch/export")
def export_evening_feedback_batch(
    class_id: int,
    payload: EveningFeedbackBatchExportRequest,
    teacher: dict = CurrentTeacher,
):
    evening_class = require_evening_class(class_id, teacher["id"])
    period = evening_period_meta(payload.period_type, payload.period_value)
    export_items = [item for item in payload.items if item.final_feedback.strip()]
    if not export_items:
        raise HTTPException(status_code=400, detail="没有可导出的晚辅反馈")

    student_ids = [item.student_id for item in export_items]
    placeholders = ",".join("?" for _ in student_ids)
    with get_db() as db:
        rows = db.execute(
            f"""
            SELECT id, name
            FROM evening_students
            WHERE teacher_id = ? AND class_id = ? AND id IN ({placeholders})
            """,
            (teacher["id"], class_id, *student_ids),
        ).fetchall()
    student_map = {row["id"]: row["name"] for row in rows}
    if len(student_map) != len(set(student_ids)):
        raise HTTPException(status_code=400, detail="导出学生不属于当前晚辅班级")
    return evening_feedback_word_response(
        evening_class=evening_class,
        period=period,
        term_label=payload.term_label,
        owner_name=payload.owner_name,
        export_subject=payload.export_subject,
        document_title=payload.document_title,
        filename_base=payload.filename_base,
        export_items=[
            {
                "student_name": item.student_name.strip() or student_map[item.student_id],
                "final_feedback": item.final_feedback.strip(),
            }
            for item in export_items
        ],
    )


@app.post("/api/evening/classes/{class_id}/feedbacks/archive/export")
def export_evening_feedback_archive(
    class_id: int,
    payload: EveningFeedbackClassExportRequest,
    teacher: dict = CurrentTeacher,
):
    evening_class = require_evening_class(class_id, teacher["id"])
    period = evening_period_meta(payload.period_type, payload.period_value)
    with get_db() as db:
        rows = db.execute(
            """
            SELECT s.name AS student_name, f.final_feedback
            FROM evening_feedbacks f
            JOIN evening_students s ON s.id = f.student_id
            WHERE f.teacher_id = ?
              AND s.teacher_id = ?
              AND s.class_id = ?
              AND f.period_type = ?
              AND f.period_value = ?
              AND TRIM(f.final_feedback) != ''
            ORDER BY s.id DESC, f.id DESC
            """,
            (teacher["id"], teacher["id"], class_id, period["period_type"], period["period_value"]),
        ).fetchall()
    export_items = [
        {"student_name": row["student_name"], "final_feedback": row["final_feedback"].strip()}
        for row in rows
    ]
    return evening_feedback_word_response(
        evening_class=evening_class,
        period=period,
        term_label=payload.term_label,
        owner_name=payload.owner_name,
        export_subject=payload.export_subject,
        document_title=payload.document_title,
        filename_base=payload.filename_base,
        export_items=export_items,
    )


@app.delete("/api/evening/feedbacks/archive/batch")
def delete_evening_feedback_archives(
    payload: EveningFeedbackArchiveDeleteRequest,
    teacher: dict = CurrentTeacher,
):
    if not payload.items:
        raise HTTPException(status_code=400, detail="请选择要删除的反馈归档")

    deleted_count = 0
    with get_db() as db:
        for item in payload.items:
            evening_class = require_evening_class(item.class_id, teacher["id"])
            period = evening_period_meta(item.period_type, item.period_value)
            cursor = db.execute(
                """
                DELETE FROM evening_feedbacks
                WHERE teacher_id = ?
                  AND period_type = ?
                  AND period_value = ?
                  AND student_id IN (
                    SELECT id FROM evening_students
                    WHERE teacher_id = ? AND class_id = ?
                  )
                """,
                (
                    teacher["id"],
                    period["period_type"],
                    period["period_value"],
                    teacher["id"],
                    evening_class["id"],
                ),
            )
            deleted_count += cursor.rowcount or 0
    return {"ok": True, "deleted_count": deleted_count}


@app.post("/api/evening/classes/{class_id}/feedbacks/batch/generate")
async def generate_evening_feedback_batch(
    class_id: int,
    payload: EveningFeedbackBatchGenerateRequest,
    teacher: dict = CurrentTeacher,
):
    require_evening_class(class_id, teacher["id"])
    period = evening_period_meta(payload.period_type, payload.period_value)
    ai_config = require_teacher_ai_config(
        teacher["id"],
        model_type=payload.model_type,
        config_id=payload.config_id,
        platform_model_id=payload.platform_model_id,
    )
    style_examples = (
        list_enabled_style_examples(teacher["id"], "evening_feedback")
        if payload.use_style_examples
        else []
    )

    async def generate_item(item) -> dict:
        result = {"student_id": item.student_id, "ok": False, "draft": "", "error": ""}
        try:
            homework_summary = item.homework_summary.strip()
            if len(homework_summary) < 5:
                raise ValueError("请至少填写 5 个字的作业完成情况")
            student = require_evening_student(item.student_id, teacher["id"])
            if student["class_id"] != class_id:
                raise ValueError("该学生不属于当前晚辅班级")
            draft = await generate_evening_feedback(
                student_name=student["name"],
                grade=student["grade"],
                school=student["school"],
                period_type=period["period_type"],
                period_label=period["period_label"],
                subject=item.subject,
                homework_summary=homework_summary,
                style_examples=style_examples,
                ai_config=ai_config,
            )
            result.update({"ok": True, "draft": draft})
        except (httpx.HTTPError, KeyError, IndexError, TypeError, ValueError) as exc:
            result["error"] = str(exc) or "AI 晚辅反馈生成失败"
        except Exception as exc:
            result["error"] = str(exc) or "AI 晚辅反馈生成失败"
        return result

    semaphore = asyncio.Semaphore(EVENING_BATCH_GENERATE_CONCURRENCY)

    async def limited_generate_item(item) -> dict:
        async with semaphore:
            return await generate_item(item)

    results = await asyncio.gather(*(limited_generate_item(item) for item in payload.items))
    return {"period": period, "results": results}


@app.post("/api/evening/classes/{class_id}/feedbacks/batch")
def save_evening_feedback_batch(
    class_id: int,
    payload: EveningFeedbackBatchSaveRequest,
    teacher: dict = CurrentTeacher,
):
    require_evening_class(class_id, teacher["id"])
    period = evening_period_meta(payload.period_type, payload.period_value)
    results = []
    timestamp = now_iso()
    with get_db() as db:
        for item in payload.items:
            result = {"student_id": item.student_id, "ok": False, "feedback": None, "error": ""}
            try:
                homework_summary = item.homework_summary.strip()
                final_feedback = item.final_feedback.strip()
                ai_draft = item.ai_draft.strip() or final_feedback
                if len(homework_summary) < 5:
                    raise ValueError("请至少填写 5 个字的作业完成情况")
                if not final_feedback:
                    raise ValueError("请先生成或填写最终反馈")
                student = require_evening_student(item.student_id, teacher["id"])
                if student["class_id"] != class_id:
                    raise ValueError("该学生不属于当前晚辅班级")

                feedback_id = item.feedback_id
                if feedback_id:
                    existing = db.execute(
                        """
                        SELECT id, student_id
                        FROM evening_feedbacks
                        WHERE id = ? AND teacher_id = ?
                        """,
                        (feedback_id, teacher["id"]),
                    ).fetchone()
                    if not existing:
                        raise ValueError("要更新的晚辅反馈不存在")
                    if existing["student_id"] != item.student_id:
                        raise ValueError("晚辅反馈和学生不匹配")
                else:
                    existing = db.execute(
                        """
                        SELECT id
                        FROM evening_feedbacks
                        WHERE teacher_id = ? AND student_id = ? AND period_type = ? AND period_value = ?
                        """,
                        (teacher["id"], item.student_id, period["period_type"], period["period_value"]),
                    ).fetchone()
                    feedback_id = existing["id"] if existing else None

                if feedback_id:
                    db.execute(
                        """
                        UPDATE evening_feedbacks
                        SET period_type = ?, period_value = ?, period_start = ?, period_end = ?,
                            period_label = ?, subject = ?, homework_summary = ?, ai_draft = ?,
                            final_feedback = ?, updated_at = ?
                        WHERE id = ? AND teacher_id = ?
                        """,
                        (
                            period["period_type"],
                            period["period_value"],
                            period["period_start"],
                            period["period_end"],
                            period["period_label"],
                            item.subject,
                            homework_summary,
                            ai_draft,
                            final_feedback,
                            timestamp,
                            feedback_id,
                            teacher["id"],
                        ),
                    )
                    saved_id = feedback_id
                else:
                    cursor = db.execute(
                        """
                        INSERT INTO evening_feedbacks (
                            teacher_id, student_id, period_type, period_value,
                            period_start, period_end, period_label, subject, homework_summary,
                            ai_draft, final_feedback, created_at, updated_at
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            teacher["id"],
                            item.student_id,
                            period["period_type"],
                            period["period_value"],
                            period["period_start"],
                            period["period_end"],
                            period["period_label"],
                            item.subject,
                            homework_summary,
                            ai_draft,
                            final_feedback,
                            timestamp,
                            timestamp,
                        ),
                    )
                    saved_id = cursor.lastrowid

                result.update(
                    {
                        "ok": True,
                        "feedback": evening_batch_feedback_row(db, saved_id, teacher["id"]),
                    }
                )
            except sqlite3.IntegrityError as exc:
                result["error"] = "该学生这个时间段已经有反馈，请刷新后再编辑"
            except ValueError as exc:
                result["error"] = str(exc)
            except Exception as exc:
                result["error"] = str(exc) or "保存失败"
            results.append(result)
    return {"period": period, "results": results}


@app.post("/api/evening/classes/{class_id}/feedbacks/generate")
async def generate_evening_draft(
    class_id: int,
    payload: EveningFeedbackGenerateRequest,
    teacher: dict = CurrentTeacher,
):
    require_evening_class(class_id, teacher["id"])
    student = require_evening_student(payload.student_id, teacher["id"])
    if student["class_id"] != class_id:
        raise HTTPException(status_code=400, detail="该学生不属于当前晚辅班级")
    period = evening_period_meta(payload.period_type, payload.period_value)
    ai_config = require_teacher_ai_config(
        teacher["id"],
        model_type=payload.model_type,
        config_id=payload.config_id,
        platform_model_id=payload.platform_model_id,
    )
    try:
        draft = await generate_evening_feedback(
            student_name=student["name"],
            grade=student["grade"],
            school=student["school"],
            period_type=period["period_type"],
            period_label=period["period_label"],
            subject=payload.subject,
            homework_summary=payload.homework_summary,
            style_examples=list_enabled_style_examples(teacher["id"], "evening_feedback")
            if payload.use_style_examples
            else [],
            ai_config=ai_config,
        )
    except (httpx.HTTPError, KeyError, IndexError, TypeError, ValueError) as exc:
        raise_ai_exception(exc, "AI 晚辅反馈生成")
    except Exception as exc:
        raise_ai_exception(exc, "AI 晚辅反馈生成")
    return {"draft": draft}


@app.post("/api/evening/classes/{class_id}/feedbacks")
def create_evening_feedback(
    class_id: int,
    payload: EveningFeedbackCreate,
    teacher: dict = CurrentTeacher,
):
    require_evening_class(class_id, teacher["id"])
    student = require_evening_student(payload.student_id, teacher["id"])
    if student["class_id"] != class_id:
        raise HTTPException(status_code=400, detail="该学生不属于当前晚辅班级")
    period = evening_period_meta(payload.period_type, payload.period_value)
    timestamp = now_iso()
    try:
        with get_db() as db:
            cursor = db.execute(
                """
                INSERT INTO evening_feedbacks (
                    teacher_id, student_id, period_type, period_value,
                    period_start, period_end, period_label, subject, homework_summary,
                    ai_draft, final_feedback, created_at, updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    teacher["id"],
                    payload.student_id,
                    period["period_type"],
                    period["period_value"],
                    period["period_start"],
                    period["period_end"],
                    period["period_label"],
                    payload.subject,
                    payload.homework_summary,
                    payload.ai_draft,
                    payload.final_feedback,
                    timestamp,
                    timestamp,
                ),
            )
            row = db.execute(
                "SELECT * FROM evening_feedbacks WHERE id = ?",
                (cursor.lastrowid,),
            ).fetchone()
    except sqlite3.IntegrityError as exc:
        raise HTTPException(status_code=400, detail="该学生这个时间段已经有反馈，请编辑已有反馈") from exc
    return {"feedback": dict(row)}


@app.put("/api/evening/feedbacks/{feedback_id}")
def update_evening_feedback(
    feedback_id: int,
    payload: EveningFeedbackUpdate,
    teacher: dict = CurrentTeacher,
):
    require_evening_feedback(feedback_id, teacher["id"])
    require_evening_student(payload.student_id, teacher["id"])
    period = evening_period_meta(payload.period_type, payload.period_value)
    try:
        with get_db() as db:
            db.execute(
                """
                UPDATE evening_feedbacks
                SET student_id = ?, period_type = ?, period_value = ?,
                    period_start = ?, period_end = ?, period_label = ?,
                    subject = ?, homework_summary = ?, ai_draft = ?, final_feedback = ?, updated_at = ?
                WHERE id = ? AND teacher_id = ?
                """,
                (
                    payload.student_id,
                    period["period_type"],
                    period["period_value"],
                    period["period_start"],
                    period["period_end"],
                    period["period_label"],
                    payload.subject,
                    payload.homework_summary,
                    payload.ai_draft,
                    payload.final_feedback,
                    now_iso(),
                    feedback_id,
                    teacher["id"],
                ),
            )
            row = db.execute(
                """
                SELECT f.*, s.name AS student_name, s.grade, s.school, c.name AS class_name
                FROM evening_feedbacks f
                JOIN evening_students s ON s.id = f.student_id
                JOIN evening_classes c ON c.id = s.class_id
                WHERE f.id = ? AND f.teacher_id = ?
                """,
                (feedback_id, teacher["id"]),
            ).fetchone()
    except sqlite3.IntegrityError as exc:
        raise HTTPException(status_code=400, detail="该学生这个时间段已经有反馈，请编辑已有反馈") from exc
    return {"feedback": dict(row)}


@app.delete("/api/evening/feedbacks/batch")
def delete_evening_feedback_batch(
    payload: EveningFeedbackDeleteRequest,
    teacher: dict = CurrentTeacher,
):
    ids = [feedback_id for feedback_id in payload.ids if feedback_id]
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要删除的晚辅反馈")
    placeholders = ",".join("?" for _ in ids)
    with get_db() as db:
        cursor = db.execute(
            f"DELETE FROM evening_feedbacks WHERE teacher_id = ? AND id IN ({placeholders})",
            (teacher["id"], *ids),
        )
    return {"ok": True, "deleted_count": cursor.rowcount or 0}


@app.delete("/api/evening/feedbacks/{feedback_id}")
def delete_evening_feedback(feedback_id: int, teacher: dict = CurrentTeacher):
    require_evening_feedback(feedback_id, teacher["id"])
    with get_db() as db:
        db.execute(
            "DELETE FROM evening_feedbacks WHERE id = ? AND teacher_id = ?",
            (feedback_id, teacher["id"]),
        )
    return {"ok": True}
